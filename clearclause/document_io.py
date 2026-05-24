"""Document loading and text normalization."""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

from .models import ParsedDocument


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_paragraphs(text: str) -> list[str]:
    blocks = re.split(r"\n\s*\n", text)
    paragraphs = [re.sub(r"\s+", " ", block).strip() for block in blocks]
    return [paragraph for paragraph in paragraphs if paragraph]


def _read_bytes(file_obj: Any) -> tuple[bytes, str]:
    if isinstance(file_obj, (str, Path)):
        path = Path(file_obj)
        return path.read_bytes(), path.name

    name = getattr(file_obj, "name", "uploaded_contract")
    if hasattr(file_obj, "seek"):
        file_obj.seek(0)
    data = file_obj.read() if hasattr(file_obj, "read") else bytes(file_obj)
    if isinstance(data, str):
        data = data.encode("utf-8")
    return data, name


def _parse_pdf(data: bytes) -> tuple[str, int, dict[str, str]]:
    import pdfplumber

    pages: list[str] = []
    metadata: dict[str, str] = {}
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        metadata = {str(k): str(v) for k, v in (pdf.metadata or {}).items()}
        for page in pdf.pages:
            pages.append(page.extract_text(x_tolerance=1, y_tolerance=3) or "")
    return "\n\n".join(pages), max(1, len(pages)), metadata


def _parse_docx(data: bytes) -> tuple[str, int, dict[str, str]]:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_cells: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            table_cells.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))

    text = "\n\n".join(paragraphs + table_cells)
    metadata = {
        "title": str(doc.core_properties.title or ""),
        "author": str(doc.core_properties.author or ""),
    }
    return text, max(1, len(paragraphs) // 18), metadata


def _parse_txt(data: bytes) -> tuple[str, int, dict[str, str]]:
    text = data.decode("utf-8", errors="replace")
    return text, max(1, len(text) // 3000), {}


def parse_document(file_obj: Any, filename: str | None = None) -> ParsedDocument:
    data, detected_name = _read_bytes(file_obj)
    name = filename or detected_name
    suffix = Path(name).suffix.lower().lstrip(".") or "txt"

    if suffix == "pdf":
        raw_text, pages, metadata = _parse_pdf(data)
        file_type = "pdf"
    elif suffix in {"docx", "doc"}:
        raw_text, pages, metadata = _parse_docx(data)
        file_type = "docx"
    else:
        raw_text, pages, metadata = _parse_txt(data)
        file_type = "txt"

    text = clean_text(raw_text)
    paragraphs = split_paragraphs(text)
    return ParsedDocument(
        filename=name,
        file_type=file_type,
        text=text,
        paragraphs=paragraphs,
        word_count=len(re.findall(r"\b\w+\b", text)),
        char_count=len(text),
        page_count=pages,
        metadata=metadata,
    )


def parse_text(text: str, filename: str = "pasted_contract.txt") -> ParsedDocument:
    cleaned = clean_text(text)
    return ParsedDocument(
        filename=filename,
        file_type="txt",
        text=cleaned,
        paragraphs=split_paragraphs(cleaned),
        word_count=len(re.findall(r"\b\w+\b", cleaned)),
        char_count=len(cleaned),
        page_count=max(1, len(cleaned) // 3000),
        metadata={},
    )

